"""Export ATS resumes to PDF, DOCX, and TXT formats."""
import os
import tempfile
import logging

log = logging.getLogger(__name__)


def _sanitize(text: str) -> str:
    replacements = {
        "\u2014": "-", "\u2013": "-", "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"', "\u2022": "*", "\u2026": "...",
        "\u00a0": " ", "\u200b": "",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def export_txt(resume_text: str, filename: str = "resume") -> str:
    path = os.path.join(tempfile.gettempdir(), f"{filename}.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(resume_text)
    return path


def export_pdf(resume_text: str, filename: str = "resume") -> str:
    from fpdf import FPDF

    resume_text = _sanitize(resume_text)
    path = os.path.join(tempfile.gettempdir(), f"{filename}.pdf")
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_margins(25, 20, 25)

    HEADINGS = {
        "PROFESSIONAL SUMMARY", "WORK EXPERIENCE", "EDUCATION",
        "SKILLS", "CERTIFICATIONS", "PROJECTS", "ACHIEVEMENTS",
        "REFERENCES", "VOLUNTEER EXPERIENCE", "LANGUAGES",
        "CAREER OBJECTIVE", "EXPERIENCE", "QUALIFICATIONS",
    }

    lines = resume_text.split("\n")
    for i, line in enumerate(lines):
        stripped = line.strip()

        if i == 0 and stripped:
            pdf.set_font("Helvetica", "B", 16)
            pdf.cell(0, 10, stripped, new_x="LMARGIN", new_y="NEXT", align="C")
            continue

        if i == 1 and ("|" in stripped or "@" in stripped):
            pdf.set_font("Helvetica", "", 9)
            pdf.cell(0, 6, stripped, new_x="LMARGIN", new_y="NEXT", align="C")
            continue

        if not stripped:
            pdf.ln(3)
            continue

        if stripped.upper() in HEADINGS or (stripped.isupper() and 3 < len(stripped) < 40):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 11)
            pdf.cell(0, 7, stripped, new_x="LMARGIN", new_y="NEXT")
            y = pdf.get_y()
            pdf.line(25, y, 185, y)
            pdf.set_xy(25, y)
            pdf.ln(2)
            continue

        if stripped.startswith(("*", "-")) and not stripped.startswith("---"):
            pdf.set_font("Helvetica", "", 10)
            bullet_text = "  * " + stripped.lstrip("*- ").strip()
            pdf.multi_cell(w=0, h=5, text=bullet_text, new_x="LMARGIN", new_y="NEXT")
            continue

        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(w=0, h=5, text=stripped, new_x="LMARGIN", new_y="NEXT")

    pdf.output(path)
    return path


def export_docx(resume_text: str, filename: str = "resume") -> str:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    HEADINGS = {
        "PROFESSIONAL SUMMARY", "WORK EXPERIENCE", "EDUCATION",
        "SKILLS", "CERTIFICATIONS", "PROJECTS", "ACHIEVEMENTS",
        "REFERENCES", "VOLUNTEER EXPERIENCE", "LANGUAGES",
        "CAREER OBJECTIVE", "EXPERIENCE", "QUALIFICATIONS",
    }

    path = os.path.join(tempfile.gettempdir(), f"{filename}.docx")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    lines = resume_text.split("\n")
    for i, line in enumerate(lines):
        stripped = line.strip()

        if i == 0 and stripped:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(16)
            continue

        if i == 1 and ("|" in stripped or "@" in stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(80, 80, 80)
            continue

        if not stripped:
            continue

        if stripped.upper() in HEADINGS or (stripped.isupper() and 3 < len(stripped) < 40):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 51, 102)
            continue

        if stripped.startswith(("*", "-")) and not stripped.startswith("---"):
            bullet_text = stripped.lstrip("*- ").strip()
            p = doc.add_paragraph(bullet_text, style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                run.font.size = Pt(10)
            continue

        p = doc.add_paragraph(stripped)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.save(path)
    return path
